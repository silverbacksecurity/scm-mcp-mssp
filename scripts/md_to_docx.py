"""Convert Markdown HLD reports to styled Word documents (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── PAN brand colours ──────────────────────────────────────────────────────────
PAN_ORANGE = RGBColor(0xFF, 0x68, 0x23)
PAN_DARK = RGBColor(0x1A, 0x1A, 0x2E)
TABLE_HEADER_BG = "FF6823"  # orange header rows
TABLE_ALT_BG = "FFF3EE"  # light orange alternate rows
LINK_BLUE = RGBColor(0x00, 0x56, 0xB3)


def _set_cell_bg(cell, hex_colour: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _add_run(
    para, text: str, bold=False, italic=False, code=False, colour: RGBColor | None = None
) -> None:
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    if colour:
        run.font.color.rgb = colour


def _parse_inline(para, text: str) -> None:
    """Render inline Markdown (bold, italic, code, links) into *para*."""
    # pattern order matters: code first, then bold+italic, bold, italic, links
    pattern = re.compile(
        r"`([^`]+)`"  # `code`
        r"|(\*\*\*|___)(.*?)\2"  # ***bold italic***
        r"|(\*\*|__)(.*?)\4"  # **bold**
        r"|(\*|_)(.*?)\6"  # *italic*
        r"|\[([^\]]+)\]\([^)]+\)"  # [text](url) — strip URL
        r"|(~~)(.*?)\9"  # ~~strikethrough~~ (render plain)
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _add_run(para, text[pos : m.start()])
        if m.group(1) is not None:  # `code`
            _add_run(para, m.group(1), code=True)
        elif m.group(3) is not None:  # ***bold italic***
            _add_run(para, m.group(3), bold=True, italic=True)
        elif m.group(5) is not None:  # **bold**
            _add_run(para, m.group(5), bold=True)
        elif m.group(7) is not None:  # *italic*
            _add_run(para, m.group(7), italic=True)
        elif m.group(8) is not None:  # [link text](url)
            _add_run(para, m.group(8), colour=LINK_BLUE)
        elif m.group(10) is not None:  # ~~strikethrough~~
            _add_run(para, m.group(10))
        pos = m.end()
    if pos < len(text):
        _add_run(para, text[pos:])


def _style_heading(para, level: int) -> None:
    if level == 1:
        para.runs[0].font.color.rgb = PAN_ORANGE
        para.runs[0].font.size = Pt(20)
    elif level == 2:
        para.runs[0].font.color.rgb = PAN_DARK
        para.runs[0].font.size = Pt(15)


def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()

    # ── Page layout ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin = section.bottom_margin = Inches(1.0)

    # ── Base paragraph style ───────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    for lvl in (1, 2, 3, 4, 5, 6):
        style_name = f"Heading {lvl}"
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = "Calibri"
            s.font.bold = True

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Blank line ─────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Heading ────────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            style = f"Heading {min(level, 4)}"
            para = doc.add_heading(level=min(level, 4))
            para.clear()
            run = para.add_run(text)
            run.bold = True
            if level == 1:
                run.font.color.rgb = PAN_ORANGE
                run.font.size = Pt(20)
            elif level == 2:
                run.font.color.rgb = PAN_DARK
                run.font.size = Pt(15)
            elif level == 3:
                run.font.size = Pt(12)
            else:
                run.font.size = Pt(10)
            i += 1
            continue

        # ── Horizontal rule ────────────────────────────────────────────────────
        if re.match(r"^[-*_]{3,}$", line.strip()):
            doc.add_paragraph("─" * 60, style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # ── Fenced code block ──────────────────────────────────────────────────
        if line.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            para = doc.add_paragraph(style="Normal")
            para.paragraph_format.left_indent = Inches(0.3)
            for cl in code_lines:
                run = para.add_run(cl + "\n")
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # ── Blockquote ─────────────────────────────────────────────────────────
        if line.startswith(">"):
            text = re.sub(r"^>+\s?", "", line)
            para = doc.add_paragraph(style="Normal")
            para.paragraph_format.left_indent = Inches(0.4)
            run = para.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Table ──────────────────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # parse rows, skip separator rows (---  :---: etc.)
            rows: list[list[str]] = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip()):
                    continue
                rows.append(cells)

            if not rows:
                continue

            ncols = max(len(r) for r in rows)
            # pad rows
            rows = [r + [""] * (ncols - len(r)) for r in rows]

            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Table Grid"

            for ri, row in enumerate(rows):
                tr = tbl.rows[ri]
                for ci, cell_text in enumerate(row):
                    cell = tr.cells[ci]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)
                    _parse_inline(para, cell_text)
                    for run in para.runs:
                        run.font.size = Pt(9)
                    if ri == 0:
                        _set_cell_bg(cell, TABLE_HEADER_BG)
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    elif ri % 2 == 0:
                        _set_cell_bg(cell, TABLE_ALT_BG)
            doc.add_paragraph()
            continue

        # ── Bullet / numbered list ─────────────────────────────────────────────
        m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if m:
            indent_chars = len(m.group(1))
            text = m.group(3)
            level = indent_chars // 2
            style = "List Bullet" if level == 0 else "List Bullet 2"
            para = doc.add_paragraph(style=style)
            para.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
            _parse_inline(para, text)
            i += 1
            continue

        # ── Meta / frontmatter lines (** key: value **) ────────────────────────
        # Bold-only lines like **Document Version:** 1.0
        if line.startswith("**") and line.count("**") >= 2:
            para = doc.add_paragraph(style="Normal")
            _parse_inline(para, line)
            i += 1
            continue

        # ── Regular paragraph ──────────────────────────────────────────────────
        para = doc.add_paragraph(style="Normal")
        _parse_inline(para, line)
        i += 1

    doc.save(out_path)


def main() -> None:
    reports_dir = Path("reports")
    out_dir = Path("reports")
    out_dir.mkdir(exist_ok=True)

    md_files = sorted(reports_dir.glob("HLD_*.md"))
    if not md_files:
        print("No HLD_*.md files found in reports/")
        sys.exit(1)

    print(f"Converting {len(md_files)} reports to Word...\n")
    errors = []
    for md in md_files:
        out = out_dir / (md.stem + ".docx")
        try:
            convert(md, out)
            size_kb = out.stat().st_size // 1024
            print(f"  ✓ {out.name} ({size_kb} KB)")
        except Exception as exc:
            import traceback

            print(f"  ✗ {md.name}: {exc}")
            traceback.print_exc()
            errors.append(md.name)

    print(f"\nDone — {len(md_files) - len(errors)} converted, {len(errors)} failed.")


if __name__ == "__main__":
    main()
