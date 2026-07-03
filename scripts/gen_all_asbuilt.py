#!/usr/bin/env python3
"""Generate as-built documents (Markdown + DOCX) for all configured tenants."""

import os
import re
import sys
import time
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from dynaconf import Dynaconf

from scm_mcp_mssp.audit.asbuilt_report import AsBuiltReportBuilder
from scm_mcp_mssp.audit.extractor import (
    extract_adem,
    extract_airs,
    extract_allocated_ips,
    extract_app_acceleration,
    extract_browser,
    extract_casb_dlp,
    extract_cdl,
    extract_enterprise_dlp,
    extract_iam_roles,
    extract_identity_sspm,
    extract_insights,
    extract_iot_security,
    extract_licenses,
    extract_mt_monitor_alerts,
    extract_ngfw_devices,
    extract_ngfw_routing,
    extract_pab_tenant,
    extract_sdwan_snapshot,
    extract_snapshot,
    extract_sspm,
    extract_traffic_steering,
    extract_ztna_connectors,
)
from scm_mcp_mssp.auth.oauth import get_scm_client
from scm_mcp_mssp.auth.sdwan import get_sdwan_client
from scm_mcp_mssp.config.settings import TenantConfig

os.environ.setdefault("SCM_MCP_MSSP_MODE", "true")

base = Dynaconf(envvar_prefix="SCM_MCP", settings_files=["settings.toml"])
secrets = Dynaconf(envvar_prefix="SCM_MCP", settings_files=[".secrets.toml"])

OUT_DIR = Path("reports/all-tenants")
OUT_DIR.mkdir(parents=True, exist_ok=True)


def md_to_docx(md: str, out_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    lines = md.split("\n")
    in_table, table_rows, in_code, code_lines = False, [], False, []
    for line in lines:
        if line.startswith("```"):
            if in_code:
                in_code = False
                if code_lines:
                    p = doc.add_paragraph("\n".join(code_lines))
                    for run in p.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(8)
                code_lines = []
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(
                set(c.replace("-", "").replace(":", "").replace(" ", "")) == set() for c in cells
            ):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_rows:
                ncols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=0, cols=ncols)
                tbl.style = "Table Grid"
                for i, row_data in enumerate(table_rows):
                    row = tbl.add_row()
                    for j, ct in enumerate(row_data[:ncols]):
                        ct = re.sub(r"\*\*(.+?)\*\*", r"\1", ct)
                        ct = re.sub(r"`(.+?)`", r"\1", ct)
                        ct = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", ct)
                        row.cells[j].text = ct
                        if i == 0:
                            for p in row.cells[j].paragraphs:
                                for run in p.runs:
                                    run.bold = True
                doc.add_paragraph()
                table_rows, in_table = [], False
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            lvl = min(len(m.group(1)), 4)
            txt = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(2))
            txt = re.sub(r"`(.+?)`", r"\1", txt)
            doc.add_heading(txt.strip(), level=lvl)
            continue
        if line.startswith(">"):
            txt = re.sub(r"\*\*(.+?)\*\*", r"\1", line.lstrip("> ").strip())
            txt = re.sub(r"`(.+?)`", r"\1", txt)
            txt = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", txt)
            if txt:
                p = doc.add_paragraph(txt)
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.italic = True
            continue
        txt = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        txt = re.sub(r"`(.+?)`", r"\1", txt)
        txt = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", txt)
        txt = re.sub(r"_(.+?)_", r"\1", txt)
        if txt.strip():
            doc.add_paragraph(txt)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


all_tenants = list(base.get("tenants").keys())
print(f"Generating as-built documents for {len(all_tenants)} tenants")
print(f"Output: {OUT_DIR.resolve()}\n")

results = []
for tname in all_tenants:
    t0 = time.time()
    try:
        cfg = dict(base.get("tenants")[tname])
        cfg.update((secrets.get("tenants") or {}).get(tname) or {})
        tc = TenantConfig(**cfg)
        label = tc.label or tname
        safe_name = re.sub(r"[^\w\-]", "-", tname.lower())
        print(f"[{tname}] Starting extraction (TSG {tc.tenant_id})...")
        sys.stdout.flush()

        client = get_scm_client(tc)
        snap = extract_snapshot(client, "Shared", str(tc.tenant_id))
        extract_licenses(client, snap)
        extract_sspm(client, snap)
        extract_identity_sspm(client, snap)
        extract_traffic_steering(client, snap)
        extract_pab_tenant(client, snap)
        extract_iam_roles(client, snap)
        extract_mt_monitor_alerts(client, snap)
        extract_adem(client, snap)
        extract_iot_security(client, snap)
        extract_casb_dlp(client, snap, "Shared")
        extract_ztna_connectors(client, snap)
        extract_browser(client, snap)
        extract_allocated_ips(client, snap)
        extract_cdl(client, snap)
        extract_insights(client, snap, region="eu")
        extract_enterprise_dlp(client, snap)
        extract_ngfw_devices(client, snap)
        extract_ngfw_routing(client, snap)
        extract_airs(client, snap)
        extract_app_acceleration(client, snap)

        # SD-WAN if available
        try:
            sdwan_client = get_sdwan_client(tc)
            extract_sdwan_snapshot(sdwan_client, snap)
        except Exception:
            pass

        elapsed_extract = time.time() - t0

        print(
            f"[{tname}] Extracted in {elapsed_extract:.0f}s | rules={len(snap.all_security_rules)} edls={len(snap.edls)} errors={len(snap.extraction_errors)}"
        )
        sys.stdout.flush()

        md = AsBuiltReportBuilder(snap, customer_name=label).to_markdown()

        md_path = OUT_DIR / f"{safe_name}-asbuilt.md"
        docx_path = OUT_DIR / f"{safe_name}-asbuilt.docx"
        md_path.write_text(md)
        md_to_docx(md, docx_path)

        elapsed = time.time() - t0
        print(
            f"[{tname}] DONE {elapsed:.0f}s | md={len(md):,}chars | docx={docx_path.stat().st_size:,}bytes"
        )
        sys.stdout.flush()
        results.append(
            {
                "tenant": tname,
                "label": label,
                "status": "OK",
                "elapsed": elapsed,
                "rules": len(snap.all_security_rules),
                "edls": len(snap.edls),
                "errors": snap.extraction_errors,
                "md": str(md_path),
                "docx": str(docx_path),
            }
        )
    except Exception as exc:
        elapsed = time.time() - t0
        print(f"[{tname}] FAILED ({elapsed:.0f}s): {exc}")
        sys.stdout.flush()
        results.append({"tenant": tname, "status": "FAIL", "elapsed": elapsed, "error": str(exc)})

print("\n" + "=" * 60)
print("SUMMARY")
print("=" * 60)
for r in results:
    if r["status"] == "OK":
        errs = f" EXTRACTION_ERRORS={len(r['errors'])}" if r["errors"] else ""
        print(f"  OK   {r['label']:<35} {r['elapsed']:.0f}s{errs}")
    else:
        print(f"  FAIL {r['tenant']:<35} {r.get('error', '')[:80]}")

ok_count = sum(1 for r in results if r["status"] == "OK")
print(f"\n{ok_count}/{len(results)} tenants completed successfully")
print(f"Files in: {OUT_DIR.resolve()}")
